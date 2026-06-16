#!/usr/bin/env python3
"""
参考资料预处理 — 全文提取工具

用法：
  python3 extract_text.py [选项] <文件路径1> [文件路径2 ...]

选项：
  -o, --output <文件>     输出到文件而不是终端

支持格式：
  .pdf → pypdf     .docx → python-docx     .xlsx/.xls → openpyxl

输出：全部内容按页打印到 stdout，可重定向或用 -o 写文件
"""

import sys, os, argparse

DEPENDENCIES = {
    ".pdf": "pypdf",
    ".docx": "python-docx",
    ".xlsx": "openpyxl",
    ".xls": "openpyxl",
}

def check_deps(path):
    ext = os.path.splitext(path)[1].lower()
    pkg = DEPENDENCIES.get(ext)
    if pkg:
        try:
            if pkg == "pypdf":
                import pypdf
            elif pkg == "python-docx":
                import docx
            elif pkg == "openpyxl":
                import openpyxl
        except ImportError:
            print(f"❌ 需要安装 {pkg}：pip3 install {pkg}", file=sys.stderr)
            sys.exit(1)

def extract_pdf(path):
    from pypdf import PdfReader
    r = PdfReader(path)
    total = len(r.pages)
    print(f"【{os.path.basename(path)}】共 {total} 页")
    empty_count = 0
    for i, page in enumerate(r.pages):
        text = page.extract_text()
        if text.strip():
            print(f"\n=== 第{i+1}页 ===")
            print(text.strip())
        else:
            empty_count += 1
    if empty_count:
        print(f"\n⚠️  {empty_count} 页无提取文本（可能是图片或扫描页），"
              f"如需读取请用 pdf-converter skill")

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    print(f"【{os.path.basename(path)}】")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for i, table in enumerate(doc.tables):
        print(f"\n--- 表格{i+1} ---")
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if any(headers):
            print(" | ".join(headers))
        for row in table.rows[1:]:
            vals = [cell.text.strip() for cell in row.cells]
            if any(v.strip() for v in vals):
                print(" | ".join(vals))

def extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    print(f"【{os.path.basename(path)}】共 {len(wb.sheetnames)} 个工作表")
    for name in wb.sheetnames:
        ws = wb[name]
        print(f"\n=== {name} ===")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) if v is not None else "" for v in row]
            if any(v.strip() for v in vals):
                print(" | ".join(vals))

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,
}

def main():
    parser = argparse.ArgumentParser(description="文档全文提取工具")
    parser.add_argument("files", nargs="+", help="文件路径")
    parser.add_argument("-o", "--output", help="输出到文件而不是终端")
    args = parser.parse_args()

    if args.output:
        sys.stdout = open(args.output, "w", encoding="utf-8")

    for path in args.files:
        if not os.path.isfile(path):
            print(f"⚠️ 文件不存在: {path}", file=sys.stderr)
            continue
        check_deps(path)
        ext = os.path.splitext(path)[1].lower()
        fn = EXTRACTORS.get(ext)
        if not fn:
            print(f"⚠️ 不支持的文件格式: {ext}", file=sys.stderr)
            continue
        try:
            fn(path)
        except Exception as e:
            print(f"❌ 读取失败 [{os.path.basename(path)}]: {e}", file=sys.stderr)

if __name__ == "__main__":
    main()
